"""
샘플 이력서 생성 - 재료공학과 학부 3학년
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 문서 생성
doc = Document()

# 제목
title = doc.add_heading('이 력 서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 인적사항
doc.add_heading('인적사항', level=1)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'

cells = [
    ('성명', '김지원'),
    ('생년월일', '2003년 03월 15일 (만 21세)'),
    ('연락처', '010-1234-5678'),
    ('이메일', 'jiwon.kim@university.ac.kr'),
    ('주소', '서울특별시 관악구 대학동')
]

for i, (label, value) in enumerate(cells):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = value

# 학력
doc.add_heading('학력', level=1)
p = doc.add_paragraph()
p.add_run('서울대학교 재료공학부\n').bold = True
p.add_run('2021.03 ~ 현재 (3학년 재학 중)\n')
p.add_run('전체 학점: 3.82 / 4.3\n')
p.add_run('전공 학점: 3.95 / 4.3')

# 주요 수강 과목
doc.add_heading('주요 수강 과목', level=1)

courses = [
    ('재료과학개론', 'A+', '재료의 구조, 결정학, 상변태 이론'),
    ('고체물리학', 'A+', '결정 구조, 밴드 이론, 반도체 물리'),
    ('반도체공정', 'A+', '증착, 식각, 리소그래피, 이온주입 등 8대 공정'),
    ('박막공학', 'A0', 'PVD, CVD, 박막 성장 메커니즘, 박막 특성'),
    ('재료분석학', 'A+', 'XRD, SEM, TEM, XPS 분석 기법'),
    ('전자재료', 'A0', '반도체, 유전체, 자성재료의 전기적 특성'),
    ('재료공정실험', 'A+', '스퍼터링, RIE, 열처리, 박막 특성 측정')
]

table = doc.add_table(rows=len(courses)+1, cols=3)
table.style = 'Light List Accent 1'

# 헤더
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '과목명'
hdr_cells[1].text = '학점'
hdr_cells[2].text = '주요 내용'

for i, (course, grade, content) in enumerate(courses, 1):
    cells = table.rows[i].cells
    cells[0].text = course
    cells[1].text = grade
    cells[2].text = content

# 프로젝트 및 연구 경험
doc.add_heading('프로젝트 및 연구 경험', level=1)

p1 = doc.add_paragraph()
p1.add_run('1. 박막 증착 공정 최적화 프로젝트\n').bold = True
p1.add_run('기간: 2024.03 ~ 2024.06 (4개월)\n')
p1.add_run('내용:\n')
p1.add_run('  • RF 스퍼터링을 이용한 ITO 박막 증착\n')
p1.add_run('  • 공정 압력, RF 파워, 기판 온도 변화에 따른 박막 특성 분석\n')
p1.add_run('  • XRD, 4-point probe, UV-Vis로 결정성, 전기전도도, 투과도 측정\n')
p1.add_run('  • 최적 조건: 3mTorr, 150W, 300℃ → 비저항 3.2×10⁻⁴ Ω·cm 달성\n')
p1.add_run('성과: 학부생 연구 우수상 수상\n')

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.add_run('2. MEMS 압력센서 제작 실습\n').bold = True
p2.add_run('기간: 2023.09 ~ 2023.12 (재료공정실험 교과)\n')
p2.add_run('내용:\n')
p2.add_run('  • Si 웨이퍼 세정 (RCA 세정)\n')
p2.add_run('  • 열산화로를 이용한 SiO₂ 박막 성장 (1000℃, 2시간)\n')
p2.add_run('  • 포토리소그래피를 이용한 패터닝 (포지티브 레지스트)\n')
p2.add_run('  • RIE를 이용한 건식 식각 (CF₄/O₂ 플라즈마)\n')
p2.add_run('  • 압력센서 다이어프램 구조 제작 및 특성 측정\n')
p2.add_run('학습: 반도체 공정의 전체 플로우 이해, 클린룸 작업 경험\n')

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.add_run('3. 저온 ALD 공정 연구 (학부생 인턴)\n').bold = True
p3.add_run('기간: 2024.07 ~ 2024.08 (여름방학, 2개월)\n')
p3.add_run('소속: 서울대 반도체공정연구실 (지도교수: 박○○ 교수님)\n')
p3.add_run('내용:\n')
p3.add_run('  • 저온 ALD를 이용한 Al₂O₃ 박막 증착 (150℃)\n')
p3.add_run('  • TMA(Trimethylaluminum)와 H₂O를 전구체로 사용\n')
p3.add_run('  • 증착 사이클, 펄스 시간에 따른 성장률 및 균일도 분석\n')
p3.add_run('  • XPS를 이용한 화학 조성 분석, AFM으로 표면 거칠기 측정\n')
p3.add_run('  • 저온에서도 우수한 박막 품질 확인 (표면 거칠기 0.3nm RMS)\n')
p3.add_run('성과: 연구실 세미나 발표, 학부생 인턴 우수상\n')

# 기술 및 장비 경험
doc.add_heading('기술 및 장비 경험', level=1)

p = doc.add_paragraph()
p.add_run('박막 증착 장비:\n').bold = True
p.add_run('  • RF/DC 스퍼터링 (PVD)\n')
p.add_run('  • 열증발 증착 (Thermal Evaporation)\n')
p.add_run('  • PECVD (Plasma Enhanced CVD)\n')
p.add_run('  • ALD (Atomic Layer Deposition)\n')
p.add_run('  • 열산화로 (Thermal Oxidation Furnace)\n')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('식각 및 패터닝:\n').bold = True
p.add_run('  • RIE (Reactive Ion Etching)\n')
p.add_run('  • 습식 식각 (HF, BOE 용액)\n')
p.add_run('  • 스핀 코터 (Spin Coater)\n')
p.add_run('  • UV 노광기 (Mask Aligner)\n')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('분석 장비:\n').bold = True
p.add_run('  • XRD (X-ray Diffraction) - 결정 구조 분석\n')
p.add_run('  • SEM (Scanning Electron Microscopy) - 표면 형상 관찰\n')
p.add_run('  • TEM (Transmission Electron Microscopy) - 기본 사용법\n')
p.add_run('  • AFM (Atomic Force Microscopy) - 표면 거칠기 측정\n')
p.add_run('  • 4-point probe - 비저항 측정\n')
p.add_run('  • UV-Vis Spectroscopy - 광학 특성 분석\n')
p.add_run('  • XPS (X-ray Photoelectron Spectroscopy) - 화학 조성 분석\n')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('소프트웨어:\n').bold = True
p.add_run('  • MATLAB - 데이터 분석 및 시뮬레이션\n')
p.add_run('  • Python - 실험 데이터 처리 (NumPy, Pandas, Matplotlib)\n')
p.add_run('  • Origin - 그래프 작성 및 분석\n')
p.add_run('  • AutoCAD - 기본 설계 도면 작성\n')

# 수상 및 자격증
doc.add_heading('수상 및 자격증', level=1)

awards = [
    ('2024.06', '학부생 연구 우수상', '서울대학교 재료공학부'),
    ('2024.08', '학부생 인턴 우수상', '서울대학교 반도체공정연구실'),
    ('2023.12', '재료공정실험 최우수 팀', '서울대학교 재료공학부'),
    ('2022.12', '학업성적 우수 장학금', '서울대학교')
]

for date, award, org in awards:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{date} - {award} ({org})')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('자격증:\n').bold = True
p.add_run('  • 컴퓨터활용능력 1급 (대한상공회의소, 2022)\n')
p.add_run('  • TOEIC Speaking Level 7 (150점, 2024)\n')

# 교내외 활동
doc.add_heading('교내외 활동', level=1)

activities = [
    ('2023.03 ~ 2024.02', '반도체 학술동아리 "나노텍" 회원', 
     '주 1회 세미나, 반도체 공정 논문 스터디, 기업 견학 참여'),
    ('2023.09 ~ 2023.12', '재료공학과 멘토링 프로그램 멘토',
     '1학년 학생 대상 전공 과목 학습 지도'),
    ('2024.05', 'SEMICON Korea 2024 참관',
     '반도체 제조 장비 전시회 참관, 산업 트렌드 파악')
]

for period, activity, detail in activities:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{period}\n').bold = True
    p.add_run(f'{activity}\n')
    p.add_run(f'{detail}')

# 관심 분야 및 향후 계획
doc.add_heading('관심 분야 및 향후 계획', level=1)

p = doc.add_paragraph()
p.add_run('관심 분야:\n').bold = True
p.add_run('  • 박막 증착 공정 (특히 ALD, CVD)\n')
p.add_run('  • 반도체 전공정 (증착, 식각, 리소그래피)\n')
p.add_run('  • 박막 특성 분석 및 공정 최적화\n')
p.add_run('  • 차세대 반도체 재료 (2D materials, High-k dielectrics)\n')

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('향후 계획:\n').bold = True
p.add_run('  • 4학년: 캡스톤 디자인 프로젝트 (반도체 공정 개선)\n')
p.add_run('  • 졸업 후: 반도체 제조 대기업 공정 엔지니어 또는 대학원 진학\n')
p.add_run('  • 장기 목표: 반도체 공정 전문가로 성장, 혁신적인 공정 기술 개발\n')

# 마지막 업데이트
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('작성일: 2024년 12월')
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 저장
doc.save('/mnt/user-data/outputs/sample_resume.docx')
print("✅ 샘플 이력서 생성 완료: sample_resume.docx")
